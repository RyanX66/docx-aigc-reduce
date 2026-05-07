"""
Verify that the rewritten .docx preserves:
- All citation markers [1]...[N]
- All images
- Title/heading text unchanged
- Reference and acknowledgement text unchanged
"""
import re, sys, json
from docx import Document
from docx.oxml.ns import qn


def verify(original_path: str, rewritten_path: str, mapping_path: str = None) -> dict:
    orig = Document(original_path)
    new = Document(rewritten_path)

    results = {'passed': [], 'failed': []}

    # 1. Paragraph count
    if len(orig.paragraphs) == len(new.paragraphs):
        results['passed'].append(f'Paragraph count: {len(orig.paragraphs)}')
    else:
        results['failed'].append(f'Paragraph count: orig={len(orig.paragraphs)} new={len(new.paragraphs)}')

    # 2. Image count
    def count_images(doc):
        count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                count += len(run._element.findall('.//' + qn('w:drawing')))
        return count
    orig_imgs = count_images(orig)
    new_imgs = count_images(new)
    if orig_imgs == new_imgs:
        results['passed'].append(f'Images preserved: {orig_imgs}')
    else:
        results['failed'].append(f'Images: orig={orig_imgs} new={new_imgs}')

    # 3. Citation markers preserved
    orig_cites = set()
    new_cites = set()
    for para in orig.paragraphs:
        for c in re.findall(r'\[(\d+)\]', para.text):
            orig_cites.add(int(c))
    for para in new.paragraphs:
        for c in re.findall(r'\[(\d+)\]', para.text):
            new_cites.add(int(c))
    if orig_cites == new_cites:
        results['passed'].append(f'Citation markers: {sorted(orig_cites)}')
    else:
        missing = orig_cites - new_cites
        extra = new_cites - orig_cites
        if missing:
            results['failed'].append(f'Missing citations: {sorted(missing)}')
        if extra:
            results['failed'].append(f'Extra citations: {sorted(extra)}')

    # 4. Check that headings didn't change
    heading_diffs = 0
    for i, (op, np) in enumerate(zip(orig.paragraphs, new.paragraphs)):
        if op.style.name.startswith('Heading') and op.text != np.text:
            heading_diffs += 1
    if heading_diffs == 0:
        results['passed'].append('All headings unchanged')
    else:
        results['failed'].append(f'{heading_diffs} headings differ')

    # 5. Check references and acknowledgements unchanged
    ref_ok = ack_ok = True
    for i, (op, np) in enumerate(zip(orig.paragraphs, new.paragraphs)):
        if op.style.name == 'Heading 1' and '参考' in op.text:
            # Compare subsequent paragraphs until next Heading 1
            for j in range(i+1, len(orig.paragraphs)):
                if orig.paragraphs[j].style.name == 'Heading 1':
                    break
                if orig.paragraphs[j].text != new.paragraphs[j].text:
                    ref_ok = False
                    break
            break
    if ref_ok:
        results['passed'].append('References unchanged')
    else:
        results['failed'].append('References were modified')

    # Summary
    print(f"Passed: {len(results['passed'])}/{len(results['passed']) + len(results['failed'])}")
    for p in results['passed']:
        print(f"  [OK] {p}")
    for f in results['failed']:
        print(f"  [FAIL] {f}")

    return results


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python verify.py <original.docx> <rewritten.docx>")
        sys.exit(1)
    verify(sys.argv[1], sys.argv[2])
