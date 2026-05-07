"""
Apply rewritten body text back to a .docx file while preserving all formatting.
Reads a JSON mapping {paragraph_index: new_text} and updates the specified paragraphs.
All other paragraphs, images, styles, and formatting are left untouched.
"""
import json, sys
from lxml import etree
from docx import Document
from docx.oxml.ns import qn


def apply_rewrite(docx_path: str, mapping_path: str, output_path: str) -> int:
    doc = Document(docx_path)

    with open(mapping_path, 'r', encoding='utf-8') as f:
        mapping = json.load(f)

    updated = 0
    for i, para in enumerate(doc.paragraphs):
        idx_str = str(i)
        if idx_str not in mapping:
            continue
        new_text = mapping[idx_str]
        runs = para.runs
        if not runs:
            continue

        # Capture formatting from first run
        first_rpr_xml = None
        rpr_elem = runs[0]._element.find(qn('w:rPr'))
        if rpr_elem is not None:
            first_rpr_xml = etree.tostring(rpr_elem, encoding='unicode')

        # Remove all existing run elements
        para_element = para._element
        for run in runs:
            para_element.remove(run._element)

        # Create new run with original formatting
        new_run = para.add_run(new_text)

        # Apply original formatting if it existed
        if first_rpr_xml is not None:
            existing_rpr = new_run._element.find(qn('w:rPr'))
            if existing_rpr is not None:
                new_run._element.remove(existing_rpr)
            new_rpr = etree.fromstring(first_rpr_xml)
            new_run._element.insert(0, new_rpr)

        updated += 1

    doc.save(output_path)
    print(f"Updated {updated} paragraphs. Saved to: {output_path}")
    return updated


if __name__ == '__main__':
    if len(sys.argv) != 4:
        print("Usage: python apply_rewrite.py <original.docx> <mapping.json> <output.docx>")
        sys.exit(1)
    apply_rewrite(sys.argv[1], sys.argv[2], sys.argv[3])
